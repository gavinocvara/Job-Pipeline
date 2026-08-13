"""
Not an LLM agent — a deterministic renderer that takes the tailored resume
JSON (already reviewed and passed) and writes a clean, ATS-safe .docx.
Kept separate from the tailoring agent on purpose: formatting should never
depend on the model getting layout right.
"""
import os
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT


def _set_font(run, size=10.5, bold=False, italic=False):
    run.font.size = Pt(size)
    run.font.name = "Calibri"
    run.bold = bold
    run.italic = italic


def build_docx(tailored_resume: dict, contact: dict, output_path: str) -> str:
    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    for m in ("top_margin", "bottom_margin", "left_margin", "right_margin"):
        setattr(section, m, Inches(0.6))

    # Header
    name = doc.add_paragraph()
    name.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = name.add_run(contact["name"])
    _set_font(r, size=16, bold=True)

    contact_line = doc.add_paragraph()
    contact_line.alignment = WD_ALIGN_PARAGRAPH.CENTER
    info = f"{contact['location']}  |  {contact['phone']}  |  {contact['email']}  |  {contact['linkedin']}"
    r = contact_line.add_run(info)
    _set_font(r, size=9.5)

    def heading(text):
        p = doc.add_paragraph()
        r = p.add_run(text.upper())
        _set_font(r, size=11, bold=True)
        p.paragraph_format.space_before = Pt(8)
        p.paragraph_format.space_after = Pt(2)
        pPr = p._p.get_or_add_pPr()
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
        pBdr = OxmlElement("w:pBdr")
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"), "single")
        bottom.set(qn("w:sz"), "6")
        bottom.set(qn("w:space"), "1")
        bottom.set(qn("w:color"), "000000")
        pBdr.append(bottom)
        pPr.append(pBdr)

    def bullet(text):
        p = doc.add_paragraph(style="List Bullet")
        r = p.add_run(text)
        _set_font(r, size=10)
        p.paragraph_format.space_after = Pt(2)

    RIGHT_TAB = Inches(7.3)

    def entry_title(title_left, title_right, sub_left="", sub_right=""):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.tab_stops.add_tab_stop(RIGHT_TAB, WD_TAB_ALIGNMENT.RIGHT)
        r = p.add_run(title_left)
        _set_font(r, size=10.5, bold=True)
        r2 = p.add_run(f"\t{title_right}")
        _set_font(r2, size=10.5, bold=True)
        if sub_left or sub_right:
            p2 = doc.add_paragraph()
            p2.paragraph_format.space_after = Pt(2)
            p2.paragraph_format.tab_stops.add_tab_stop(RIGHT_TAB, WD_TAB_ALIGNMENT.RIGHT)
            r3 = p2.add_run(sub_left)
            _set_font(r3, size=10, italic=True)
            r4 = p2.add_run(f"\t{sub_right}")
            _set_font(r4, size=10, italic=True)

    # Summary
    if tailored_resume.get("summary"):
        heading("Summary")
        p = doc.add_paragraph()
        r = p.add_run(tailored_resume["summary"])
        _set_font(r, size=10)

    # Education
    edu = tailored_resume.get("education", {})
    heading("Education")
    entry_title(edu.get("school", ""), edu.get("location", ""),
                edu.get("degree", ""), edu.get("graduation", ""))

    # Skills
    if tailored_resume.get("skills_to_show"):
        heading("Technical Skills")
        p = doc.add_paragraph()
        r = p.add_run(", ".join(tailored_resume["skills_to_show"]))
        _set_font(r, size=10)

    # Experience
    if tailored_resume.get("experience"):
        heading("Experience")
        for job in tailored_resume["experience"]:
            entry_title(job.get("org", ""), job.get("location", ""),
                        job.get("title", ""), job.get("dates", ""))
            for b in job.get("bullets", []):
                bullet(b)

    # Projects
    if tailored_resume.get("projects"):
        heading("Projects")
        for proj in tailored_resume["projects"]:
            entry_title(proj.get("name", ""), proj.get("location", ""),
                        proj.get("role", ""), proj.get("dates", ""))
            for b in proj.get("bullets", []):
                bullet(b)

    os.makedirs(os.path.dirname(output_path), exist_ok=True)
    doc.save(output_path)
    return output_path
